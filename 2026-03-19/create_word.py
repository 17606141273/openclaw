from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

doc = Document()

# 设置中文字体
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# 标题
title = doc.add_heading('徐宁波 | 软件测试工程师', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息
info = doc.add_paragraph()
info.add_run('电话：').bold = True
info.add_run('17633887100  |  ')
info.add_run('邮箱：').bold = True
info.add_run('bobo1273@163.com  |  ')
info.add_run('现居：').bold = True
info.add_run('长沙市  |  ')
info.add_run('求职意向：').bold = True
info.add_run('杭州 · 软件测试/自动化测试工程师')

# 工作经历
doc.add_heading('工作经历', level=1)

doc.add_heading('杭州七凌科技有限公司（字节跳动/抖音项目）｜软件测试工程师', level=2)
p = doc.add_paragraph()
p.add_run('2024.7 – 至今').italic = True
doc.add_paragraph('负责抖音增长业务（青少年、投放承接、桌面组件）质量保障，支撑产品快速迭代与稳定上线')
doc.add_paragraph('基于Tesla平台搭建接口自动化测试体系，编写并维护用例库，替代80%重复手动回归工作')
doc.add_paragraph('使用Shoots+Python实现抖音及同构版整体投放链路UI自动化，覆盖Android/iOS全版本，用例通过率稳定在95%+')
doc.add_paragraph('基于Kepler平台完成全链路监控接入与异常检测，精准识别大促场景服务异常')
doc.add_paragraph('通过混沌平台模拟极端场景，验证系统容错与故障恢复能力')
doc.add_paragraph('规范Bug记录与生命周期管理，推动Bug高效闭环')

doc.add_heading('博彦科技股份有限公司（Automation Team）｜自动化测试工程师', level=2)
p = doc.add_paragraph()
p.add_run('2021.7 – 2024.3').italic = True
doc.add_paragraph('主导设计并落地基于Python+Selenium的Web端自动化测试框架，核心测试覆盖率85%+')
doc.add_paragraph('基于Jenkins搭建CI/CD自动化测试流水线，实现测试脚本自动触发、执行及报告自动输出')
doc.add_paragraph('独立搭建性能测试体系，运用JMeter模拟高并发场景，使平台高峰期响应时间缩短50%')
doc.add_paragraph('使用Jira进行缺陷全生命周期管理，推动问题高效闭环')

# 专业技能
doc.add_heading('专业技能', level=1)

skills = [
    ('自动化测试', 'Selenium + Python + Pytest + Yaml + POM 搭建Web UI自动化框架'),
    ('接口测试', 'Requests + Pytest + Allure 接口自动化，Tesla 框架全流程落地'),
    ('性能测试', 'JMeter 脚本创建、场景设置、报告分析及瓶颈定位'),
    ('持续集成', 'Jenkins 持续集成、自动部署、定时任务配置'),
    ('测试工具', '禅道、Meego、Jira、Git、MySQL、Linux'),
    ('平台经验', '字节跳动 Kepler 智能运维平台、Shoots、BitS、混沌平台'),
]

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '类别'
hdr[1].text = '技能'
for i, (cat, skill) in enumerate(skills):
    row = table.rows[i+1].cells
    row[0].text = cat
    row[1].text = skill

# 项目经历
doc.add_heading('项目经历', level=1)

doc.add_heading('抖音UG投放拉活用户厂商位置定制策略｜项目QA', level=2)
p = doc.add_paragraph()
p.add_run('2025.7-2025.10').italic = True
doc.add_paragraph('深度解读需求文档与接口文档，拆解厂商适配规则与用户分层触达条件')
doc.add_paragraph('主导多厂商手机（华为、小米、OPPO、vivo等）及不同系统版本的兼容性测试')
doc.add_paragraph('基于Tesla平台编写接口自动化用例，与功能测试同步进行，测试效率提升60%')
doc.add_paragraph('利用GUI智能测试平台验证埋点数据，缩短70%测试时间')
doc.add_paragraph('编写自动化脚本覆盖约80%核心回归用例，提升集成回归效率')

doc.add_heading('抖音UG投放承接项目｜项目QA', level=2)
p = doc.add_paragraph()
p.add_run('2024.7 – 至今').italic = True
doc.add_paragraph('负责用户从端外广告点击到进端的整体流程测试，保障"所见即所得"')
doc.add_paragraph('基于Tesla平台实现接口自动化，高效消费测试结果')
doc.add_paragraph('利用GUI智能测试平台对庞大埋点数据进行验证，提前规避多处问题')
doc.add_paragraph('通过BitS平台设置每周定时巡检，避免版本迭代引入回归Bug')

# 教育背景
doc.add_heading('教育背景', level=1)
doc.add_paragraph('宿迁职业技术学院 ｜ 移动应用开发 ｜ 专科（2020.10 – 2023.7）')

# 自我评价
doc.add_heading('自我评价', level=1)
doc.add_paragraph('具备4年互联网软件测试实战经验，熟悉敏捷迭代流程，精通功能、兼容性、接口、性能及自动化全维度测试。擅长自动化框架搭建与CI/CD落地，注重测试效率与质量保障，能快速定位问题并沉淀测试方案，以严谨的测试思维保障产品稳定上线与迭代优化。')

# 保存Word文档
doc.save(r'D:\AI_Files\2026-03-19\徐宁波-简历-优化版.docx')
print('Word文档创建成功！')
