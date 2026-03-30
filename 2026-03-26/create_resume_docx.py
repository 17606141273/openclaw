# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# 标题
title = doc.add_heading('徐宁波', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 头部信息
info_para = doc.add_paragraph()
info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = info_para.add_run('软件测试/自动化测试工程师  |  杭州  |  17633887100  |  bobo1273@163.com  |  25岁')
run.font.size = Pt(10)

# 专业技能
doc.add_heading('专业技能', level=1)
skills = [
    '熟练使用 Selenium+Python+Pytest+YAML+POM 搭建 WebUI 自动化测试框架，配合 Jenkins 实现持续集成与定时任务',
    '掌握 UFT (QTP) 自动化测试工具，能够独立完成 Web/桌面端功能自动化脚本开发与维护',
    '精通 Python 接口自动化测试，熟练使用 Requests 库构建接口请求、Pytest 管理测试用例，结合 Allure 生成可视化报告',
    '掌握 Tesla 接口自动化框架（字节内部）全流程落地，擅长构建可复用物料库',
    '熟练使用 Shoots+Python（字节内部）完成抖音及同构版整体投放链路（端外广告点击+进端后目标落地页各种结果验证及操作）自动化用例开发，覆盖 Android/iOS 全版本，通过率稳定在 95%+，有效拦截 80%+ 机型适配问题',
    '熟悉性能测试流程，熟练使用 JMeter 工具进行创建脚本、场景设置及报告生成，对核心指标进行分析，曾参与项目瓶颈定位与调优',
    '熟练使用字节跳动 Kepler 智能运维平台（字节内部），可独立完成全链路监控接入与智能异常检测配置，精准识别版本发布及大促场景服务异常，借助多源数据关联快速定位根因',
    '精通业务级容灾测试体系搭建，熟练使用 Chaos 混沌平台（字节内部）设计流量突增、依赖服务宕机、数据异常等极端场景，验证系统容错与故障恢复能力',
    '熟练掌握 Linux 常用命令及 MySQL 数据库操作，熟悉禅道，Jira 等缺陷管理工具'
]
for skill in skills:
    p = doc.add_paragraph(skill, style='List Bullet')

# 工作经历
doc.add_heading('工作经历', level=1)

doc.add_heading('杭州七凌科技有限公司 | 软件测试工程师', level=2)
doc.add_paragraph('字节跳动/抖音项目 | 2024.7 - 至今')
doc.add_paragraph('负责抖音增长项目组投放承接业务线，游戏、社交、电商三大场景端外广告投放链路测试工作，聚焦投放业务拉新、拉活核心目标，重点验证广告点击跳转、参数透传、落地页渲染展示全流程功能正确性，严格把控端外广告与落地页内容一致性，实现用户所见即所得，杜绝信息不符影响拉新拉活转化的问题。参与需求评审、测试用例设计、多场景交叉测试、缺陷定位跟踪及回归验证，推动问题全流程闭环，保障各场景投放链路稳定顺畅，为投放端拉新拉活业务的高效落地、用户体验优化提供质量保障，助力业务增长目标达成。')

jobs1 = [
    '负责抖音增长业务需求测试，聚焦核心增长场景，保障业务快速迭代与稳定上线，支撑产品用户规模与营收增长目标达成',
    '主导 Tesla 接口自动化框架落地，编写维护 500+ 用例，测试效率提升 60%',
    '使用 Shoots+Python 实现抖音投放链路 UI 自动化，覆盖 Android/iOS 全版本，通过率 95%+，有效拦截 80%+ 机型适配问题',
    '基于 Kepler 智能运维平台搭建全链路监控，精准识别大促场景服务异常，借助多源数据关联快速定位根因',
    '通过 Chaos 混沌平台开展业务容灾测试，设计流量突增、依赖服务宕机等极端场景，验证系统容错与故障恢复能力',
    '基于 Jenkins 搭建 CI/CD 自动化测试流水线，实现测试脚本自动触发与报告输出',
    '输出测试计划，用例、报告等规范化文档，量化测试覆盖度、缺陷密度等核心指标'
]
for job in jobs1:
    doc.add_paragraph(job, style='List Bullet')

doc.add_heading('博彦科技股份有限公司 | 自动化测试工程师', level=2)
doc.add_paragraph('Automation team/Accipiens7-8 | 2021.7 - 2024.3')
jobs2 = [
    '基于 Python + QTP/UFT 搭建 Web/桌面端自动化测试框架，核心覆盖率 85%+',
    '基于 Jenkins 搭建 CI/CD 自动化测试流水线，实现测试脚本自动触发与报告输出',
    '运用 JMeter 开展性能测试、压力测试，高峰期响应时间缩短 50%',
    '使用 Jira 进行缺陷全生命周期管理，推动 Bug 高效闭环'
]
for job in jobs2:
    doc.add_paragraph(job, style='List Bullet')

# 项目经历
doc.add_heading('项目经历', level=1)

doc.add_heading('抖音 UG 投放承接 - 小游戏承接体裁尝试', level=2)
doc.add_paragraph('投放拉活用户进端后，用弹窗、snackbar形式承接，对比线上小程序/锚点视频承接，探索用户留存最优承接')
p1 = [
    '深度解读需求文档，重点拆解抖音弹窗管理器优先级，制定覆盖功能，安全、边界、异常场景的测试方案',
    '验证所投放广告素材中不包含对应的游戏id及包含多个id时三种弹窗可正常展示及交互，确保进端后展示的小游戏与广告素材一致',
    '利用 Shoots+Python 编写核心测试用例并在开发提测演示时同步执行冒烟测试，加快整体测试进度',
    '模拟极端场景，验证系统容错能力与异常提示清晰度，测试新旧方案切换兼容性',
    '利用 GUI 智能测试平台编写 AI 用例脚本验证埋点数据，缩短测试时间',
    '主导多厂商主流机型兼容性测试，重点验证弹窗UI界面展示是否出现异变或不可交互问题',
    '利用 Shoots+Python 编写 UI 自动化脚本，覆盖 60% 核心回归用例，提升回归效率',
    '规范记录 Bug 细节，精准判定严重级/优先级，协同开发定位问题并推动修复，确保上线前关键问题 100%修复',
    '使用 Kepler 平台完成全链路监控接入与智能异常检测配置',
    '输出标准化测试报告，沉淀投放游戏测试经验，形成可复用测试用例库'
]
for item in p1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('抖音 UG 投放承接 - 合养精灵承接弹窗使用喂食元素', level=2)
doc.add_paragraph('投放【合养精灵】素材的用户，进入抖音/抖极后，承接弹窗异化。尝试用发皮肤+喂养美食的利益点吸引用户')
p2 = [
    '深度解读需求文档，制定覆盖功能，安全、边界、异常场景的测试方案',
    '利用 GUI 智能测试平台编写 AI 用例脚本进行新需求测试并编写埋点用例进行验证',
    'Tesla 平台接口自动化编写，覆盖核心回归场景',
    '上线后将 GUI 测试用例使用 Bits 平台设置定时巡检，提升回归效率',
    '规范记录 Bug 细节，精准判定严重级/优先级，协同开发定位问题并推动修复，确保上线前关键问题 100%修复',
    '使用 Kepler 平台完成全链路监控接入与智能异常检测配置',
    '输出标准化测试报告，沉淀投放社交场景测试经验，形成可复用测试用例库'
]
for item in p2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('抖音 UG 投放承接 - 秒杀页灌入联投商品', level=2)
doc.add_paragraph('探索性需求主要是尝试运用秒杀页面+灌入联投商品，看能否在该页面提高转化，提升联投商品交易总额收益')
p3 = [
    '深度解读需求文档，重点拆解抖音秒杀七分屏的展示逻辑及功能点，制定覆盖功能，安全、边界、异常场景的测试方案',
    '验证所投放广告素材中不包含对应的商品id及包含多个商品id时对应秒杀页展示是否正常，确保进端后展示的置顶商品与广告素材一致',
    '利用 GUI 智能测试平台编写 AI 用例脚本进行新需求测试并编写埋点用例进行验证',
    'Tesla 平台接口自动化编写，覆盖核心回归场景',
    '上线后将 GUI 测试用例使用 Bits 平台设置定时巡检，提升回归效率',
    '规范记录 Bug 细节，精准判定严重级/优先级，协同开发定位问题并推动修复，确保上线前关键问题 100%修复',
    '使用 Kepler 平台完成全链路监控接入与智能异常检测配置',
    '输出标准化测试报告，沉淀投放电商场景测试经验，形成可复用测试用例库'
]
for item in p3:
    doc.add_paragraph(item, style='List Bullet')

# 教育经历
doc.add_heading('教育经历', level=1)
doc.add_paragraph('国家开放大学 | 计算机科学与技术 | 本科')
doc.add_paragraph('宿迁职业技术学院 | 移动应用开发 | 专科')

# 自我评价
doc.add_heading('自我评价', level=1)
doc.add_paragraph('具备多年互联网软件测试实战经验，熟悉互联网产品敏捷迭代流程，精通功能、兼容性、接口、性能及自动化全维度测试。擅长自动化测试框架搭建、CI/CD 流水线落地与多厂商机型适配测试，注重测试效率与质量保障，能快速定位问题、沉淀测试方案，以严谨的测试思维保障产品稳定上线与迭代优化。')

doc.save('D:/AI_Files/2026-03-26/徐宁波-软件测试工程师-简历.docx')
print('Done')
