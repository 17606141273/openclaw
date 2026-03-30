from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

# Read markdown
with open('D:/AI_Files/2026-03-21/徐宁波-软件测试工程师-简历-优化版.md', 'r', encoding='utf-8') as f:
    md = f.read()

html = markdown.markdown(md)

doc = Document()

# Title
title = doc.add_heading('徐宁波 | 软件测试工程师', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact info
p = doc.add_paragraph()
p.add_run('📞 17633887100  |  📧 bobo1273@163.com  |  📍 杭州市  |  🎂 25岁').font.size = Pt(11)

doc.add_paragraph()

# Process content
lines = md.split('\n')
for line in lines:
    if line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- '):
        doc.add_paragraph(line, style='List Bullet')
    elif line.strip():
        doc.add_paragraph(line)

doc.save('D:/AI_Files/2026-03-21/徐宁波-简历-优化版.docx')
print("Word doc created!")
