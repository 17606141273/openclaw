from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

# Title
title = doc.add_heading('徐宁波 | 软件测试工程师', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact info
contact = doc.add_paragraph()
contact.add_run('📞 17633887100  |  📧 bobo1273@163.com  |  📍 杭州市  |  🎂 25岁').font.size = Pt(10)
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Job intention
doc.add_heading('求职意向', level=1)
p = doc.add_paragraph()
p.add_run('岗位：').bold = True
p.add_run('自动化测试工程师 / 软件测试工程师')
p = doc.add_paragraph()
p.add_run('薪资：').bold = True
p.add_run('面议')
p = doc.add_paragraph()
p.add_run('类型：').bold = True
p.add_run('社招')

# Skills
doc.add_heading('专业技能', level=1)

# Automation testing
doc.add_heading('自动化测试', level=2)
doc.add_paragraph('精通 Selenium + Python + Pytest + YAML + POM 搭建 Web UI 自动化测试框架，配置文件管理环境，适配企业需求', style='List Bullet')
doc.add_paragraph('熟练使用 Shoots + Python 完成抖音/iOS/Android 端外广告投放链路自动化，自动化用例通过率 95%+，拦截 80%+ 机型适配问题', style='List Bullet')
doc.add_paragraph('精通 Python 接口自动化测试（Requests + Pytest + Allure），覆盖正向/反向/边界/异常场景', style='List Bullet')

# CI/CD
doc.add_heading('持续集成', level=2)
doc.add_paragraph('熟练使用 Jenkins 实现持续集成、自动部署和定时任务，独立搭建并落地 Accipiens7-8 项目 CI/CD 流水线', style='List Bullet')

# Performance
doc.add_heading('性能测试', level=2)
doc.add_paragraph('熟悉性能测试流程，熟练使用 JMeter 进行性能测试、瓶颈定位与调优，平台高峰期响应时间缩短 50%', style='List Bullet')

# Other tools
doc.add_heading('其他工具', level=2)
doc.add_paragraph('字节跳动 Kepler 智能运维平台、混沌平台、Meego/禅道、Linux、MySQL', style='List Bullet')

# Work experience
doc.add_heading('工作经历', level=1)

# Company 1
doc.add_heading('杭州七凌科技有限公司 | 软件测试工程师', level=2)
p = doc.add_paragraph()
p.add_run('2024.07 - 至今  |  字节跳动/抖音项目').italic = True

responsibilities = [
    ('业务测试与增长保障', '负责抖音增长业务（青少年、投放承接、桌面组件）测试，保障业务快速迭代与稳定上线'),
    ('自动化测试体系搭建', '使用 Tesla 平台实现接口自动化测试，编写并维护自动化用例库\n使用 Shoots + Python 实现 UI 自动化落地，替代重复手动回归工作'),
    ('智能运维与容灾验证', '基于 Kepler 平台完成全链路监控接入与智能异常检测配置\n通过混沌平台开展业务容灾测试，验证系统容错与恢复能力'),
    ('测试计划与缺陷管理', '拆解需求，制定覆盖功能/边界/异常的测试计划，设计高复用性测试用例\n规范缺陷管理，推动 Bug 闭环，确保上线前关键问题 100% 修复')
]

for title, desc in responsibilities:
    p = doc.add_paragraph(style='List Number')
    p.add_run(title).bold = True
    for line in desc.split('\n'):
        doc.add_paragraph(line, style='List Bullet 2')

# Company 2
doc.add_heading('博彦科技股份有限公司 | 自动化测试工程师', level=2)
p = doc.add_paragraph()
p.add_run('2021.07 - 2024.03  |  Automation Team / Accipiens7-8').italic = True

responsibilities2 = [
    ('自动化测试框架搭建', '基于 Python + QTP/UFT 搭建 Web/桌面端功能自动化测试体系\n核心测试覆盖率稳定保持 85%+'),
    ('CI/CD 流水线', '基于 Jenkins 搭建自动化测试流水线，实现测试脚本自动触发，执行及报告输出'),
    ('性能测试', '独立搭建性能测试体系，运用 JMeter 模拟高并发场景\n定位核心性能瓶颈并输出优化方案，平台响应时间缩短 50%')
]

for title, desc in responsibilities2:
    p = doc.add_paragraph(style='List Number')
    p.add_run(title).bold = True
    for line in desc.split('\n'):
        doc.add_paragraph(line, style='List Bullet 2')

# Projects
doc.add_heading('项目经历', level=1)

# Project 1
doc.add_heading('抖音UG投放拉活用户厂商位置定制策略 | 项目QA', level=2)
p = doc.add_paragraph()
p.add_run('2025.07 - 2025.10').italic = True

project1_points = [
    '深度解读需求，制定覆盖功能/兼容性/边界/异常的测试方案',
    '主导多厂商、多系统版本兼容性测试，验证厂商免费位置展示效果',
    '在 Tesla 平台编写接口自动化用例，与功能测试同步进行，提高测试效率',
    '利用 GUI 智能测试平台编写 AI 用例验证埋点数据，缩短测试时间',
    '输出标准化测试报告，量化测试覆盖度、缺陷密度等核心指标'
]
for point in project1_points:
    doc.add_paragraph(point, style='List Bullet')

# Project 2
doc.add_heading('抖音UG投放承接项目 | 项目QA', level=2)
p = doc.add_paragraph()
p.add_run('2024.07 - 至今').italic = True

project2_points = [
    '负责用户进端后所见、用户分层触达条件的测试方案设计',
    '在 Tesla 平台编写接口自动化用例，提高测试效率',
    '利用 GUI 平台编写自动化脚本，覆盖 80% 核心回归用例',
    '每周定时巡检执行，避免版本迭代引入回归 Bug'
]
for point in project2_points:
    doc.add_paragraph(point, style='List Bullet')

# Education
doc.add_heading('教育经历', level=1)
p = doc.add_paragraph()
p.add_run('宿迁职业技术学院 | 移动应用开发 | 2020.10 - 2023.07').bold = True
p.add_run('（专科）')

# Self evaluation
doc.add_heading('自我评价', level=1)
eval_points = [
    '具备多年互联网软件测试实战经验，熟悉敏捷迭代流程',
    '精通功能、兼容性、接口、性能及自动化全维度测试',
    '擅长自动化测试框架搭建与 CI/CD 流水线落地',
    '注重测试效率与质量保障，能快速定位问题',
    '具备良好的团队协作与沟通能力，善于推动问题闭环'
]
for point in eval_points:
    doc.add_paragraph(point, style='List Bullet')

# Save
doc.save('D:/AI_Files/2026-03-21/徐宁波-简历-优化版.docx')
print("Word document created successfully!")
